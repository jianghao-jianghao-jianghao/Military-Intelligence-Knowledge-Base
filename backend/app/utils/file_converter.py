
import io
import re
import docx
from fastapi import UploadFile
from docx.enum.text import WD_ALIGN_PARAGRAPH
from bs4 import BeautifulSoup
from app.core.parsers.factory import parser_factory

# 注意：生产环境建议使用 WeasyPrint 或 Playwright 进行高质量 PDF 生成
try:
    from weasyprint import HTML, CSS
except ImportError:
    HTML = None

try:
    from htmldocx import HtmlToDocx
except ImportError:
    HtmlToDocx = None

class FileConverter:
    """
    文件格式转换工具类 (Facade Pattern)
    
    职责：
    1. 导入 (Ingestion): 委托给 `app.core.parsers` 处理，将各种格式转为 Markdown。
    2. 导出 (Export): 将 HTML (LLM 生成结果) 转换为 PDF/Word 供下载。
    3. 分析 (Analysis): 专门针对 Docx 的样式分析逻辑 (智能校对用)。
    """

    async def parse_to_markdown(self, file: UploadFile) -> str:
        """
        [重构后] 通用入口：将上传的文件内容解析为 Markdown 格式字符串
        现在的逻辑非常清晰：直接找工厂要解析器，然后执行。
        """
        parser = parser_factory.get_parser(file.filename)
        return await parser.parse(file)

    async def analyze_document_style(self, file: UploadFile) -> str:
        """
        [智能校对核心]
        分析参考文档 (Docx/Dot) 的格式特征，并生成 Markdown 格式的“样式规范描述”。
        LLM 将基于此描述来检查 Draft 是否符合规范。
        
        注：此逻辑与“文本提取”不同，属于特定的“样式提取”，因此保留在此处，或未来移入专门的 StyleAnalyzer。
        """
        filename = file.filename.lower()
        if not (filename.endswith(".docx") or filename.endswith(".dotx") or filename.endswith(".dot")):
            return "仅支持 Word (.docx, .dotx) 格式的样式学习。"

        content = await file.read()
        file.file.seek(0)
        
        try:
            doc = docx.Document(io.BytesIO(content))
            style_report = []
            style_report.append(f"# 学习到的文档格式规范 ({filename})")
            
            # 1. 分析页面设置 (Page Setup)
            # 注意：python-docx 读取 section 属性
            if doc.sections:
                section = doc.sections[0]
                # 简易转换 EMU 到 厘米/英寸
                # 这里仅做定性描述
                style_report.append("\n## 1. 页面版式")
                style_report.append("- **纸张方向**: " + ("横向" if section.orientation == 1 else "纵向"))
                # style_report.append(f"- **页边距**: 上下左右约 {int(section.top_margin.cm)}cm")

            # 2. 分析段落样式 (Paragraph Styles)
            style_report.append("\n## 2. 段落与字体规范")
            
            # 统计常用样式的特征
            # 这是一个启发式算法：采样前 50 个段落，归纳特征
            seen_styles = set()
            
            for para in doc.paragraphs[:50]:
                if not para.text.strip():
                    continue
                
                style_name = para.style.name
                if style_name in seen_styles:
                    continue
                seen_styles.add(style_name)
                
                # 提取特征
                alignment_map = {
                    WD_ALIGN_PARAGRAPH.LEFT: "左对齐",
                    WD_ALIGN_PARAGRAPH.CENTER: "居中",
                    WD_ALIGN_PARAGRAPH.RIGHT: "右对齐",
                    WD_ALIGN_PARAGRAPH.JUSTIFY: "两端对齐"
                }
                align = alignment_map.get(para.alignment, "默认对齐")
                
                # 尝试获取字体名称 (python-docx 获取字体比较繁琐，需检查 run 或 style 继承)
                font_name = "默认字体"
                font_size = "默认字号"
                if para.runs:
                    run = para.runs[0]
                    if run.font.name:
                        font_name = run.font.name
                    elif run.font.name_far_east: # 中文字体
                        font_name = run.font.name_far_east
                    
                    if run.font.size:
                        # Pt 转换
                        font_size = f"{int(run.font.size.pt)}pt"

                # 识别是否为标题
                role = "正文"
                if "Heading" in style_name or "标题" in style_name:
                    role = "标题"
                
                desc = f"- **{style_name}** ({role}): {align}, 字体[{font_name}], 字号[{font_size}]"
                if "Red" in style_name or "红头" in style_name:
                    desc += " (检测到红头特征)"
                
                style_report.append(desc)

            # 3. 提取特定关键词规则 (Content Rules)
            style_report.append("\n## 3. 内容要素要求")
            full_text = "\n".join([p.text for p in doc.paragraphs])
            
            if "签发人" in full_text:
                style_report.append("- 必须包含“签发人”及其对应的姓名。")
            if "主题词" in full_text:
                style_report.append("- 文末必须包含“主题词”区域。")
            if re.search(r"二[〇○]二[一二三四]年", full_text):
                style_report.append("- 日期必须使用汉字数字格式 (如：二〇二四年)。")
            
            return "\n".join(style_report)

        except Exception as e:
            return f"样式解析失败: {str(e)}"

    # --- 导出功能 (Export) 保持不变，也可视情况重构为 Exporter 策略 ---

    def html_to_pdf(self, html_content: str) -> bytes:
        """
        将 HTML 转换为 PDF 字节流 (A4 标准)
        """
        if not HTML:
            raise NotImplementedError("WeasyPrint library not installed.")
        
        # 注入基础打印样式，确保 PDF 渲染精确
        base_css = CSS(string="""
            @page { size: A4; margin: 0; }
            body { font-family: 'SimSun', serif; }
        """)
        
        pdf_bytes = HTML(string=html_content).write_pdf(stylesheets=[base_css])
        return pdf_bytes

    def html_to_docx(self, html_content: str) -> io.BytesIO:
        """
        将 HTML 转换为 Word 文档流
        """
        if not HtmlToDocx:
            raise NotImplementedError("htmldocx library not installed.")

        # 清理 HTML，移除对 Word 不友好的标签
        soup = BeautifulSoup(html_content, 'html.parser')
        # 移除 scripts, styles
        for script in soup(["script", "style"]):
            script.decompose()
            
        clean_html = str(soup)
        
        new_parser = HtmlToDocx()
        doc = docx.Document()
        new_parser.add_html_to_document(clean_html, doc)
        
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        return output

file_converter = FileConverter()
